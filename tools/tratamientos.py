from docx import Document

def text_to_word(text, output_filename="output.docx"):
    doc = Document()
    lines = text.split("\n")
    
    for line in lines:
        stripped = line.strip()

        if stripped.startswith("# "):  # Encabezado nivel 1
            doc.add_heading(stripped[2:].replace("**", ""), level=1)
        elif stripped.startswith("## "):  # Encabezado nivel 2
            doc.add_heading(stripped[3:].replace("**", ""), level=2)
        elif stripped.startswith("### "):  # Encabezado nivel 3
            doc.add_heading(stripped[4:].replace("**", ""), level=3)
        elif stripped.startswith("- "):  # Lista
            doc.add_paragraph(stripped[2:].replace("**", ""), style="List Bullet")
        elif stripped.startswith("**") and stripped.endswith("**"):  # Negrita
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(stripped[2:-2])
            run.bold = True
        elif stripped.startswith("```"):  # Bloque de código
            continue  # Podrías agregar soporte para bloques si lo deseas
        elif stripped:
            doc.add_paragraph(stripped.replace("**", ""))
    
    doc.save(output_filename)
    print(f"✅ Documento guardado como {output_filename}")